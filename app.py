import streamlit as st
import pandas as pd
import random
import json
import re
import requests
import os
import docx
from datetime import datetime, timedelta
from faker import Faker
from google import genai
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document

st.set_page_config(
    page_title="Chatbot INAPROC - AI Assistant",
    page_icon="🤖",
    layout="wide"
)

st.markdown("""
<style>
    :root {
        --merah-inaproc: #B22222;
        --merah-muda: #FFE4E1;
        --merah-tua: #8B0000;
    }
    .stButton > button {
        background-color: var(--merah-inaproc);
        color: white;
        border: none;
        border-radius: 8px;
        width: 100%;
    }
    .stButton > button:hover {
        background-color: var(--merah-tua);
    }
    .stChatInput input {
        border: 2px solid var(--merah-inaproc);
        border-radius: 20px;
    }
    [data-testid="stSidebar"] {
        background-color: #FFF8F5;
        border-right: 3px solid var(--merah-inaproc);
    }
    .user-profile-card {
        background-color: #FFE4E1;
        padding: 15px;
        border-radius: 10px;
        margin-bottom: 20px;
        text-align: center;
    }
    .role-badge {
        background-color: #B22222;
        color: white;
        padding: 5px 10px;
        border-radius: 20px;
        font-size: 12px;
        display: inline-block;
    }
</style>
""", unsafe_allow_html=True)

try:
    API_KEY = st.secrets["google_api_key"]
    client = genai.Client(api_key=API_KEY)
except Exception as e:
    st.error(f"❌ Gagal mengambil API Key: {e}")
    st.stop()

USER_CREDENTIALS = {
    "ppk_budi": {"password": "12345", "nama": "Budi Santoso", "role": "PPK", "unit": "Dinas Pendidikan"},
    "staff_siti": {"password": "12345", "nama": "Siti Aminah", "role": "Staff Pengadaan", "unit": "Dinas Kesehatan"},
    "vendor_ahmad": {"password": "12345", "nama": "Ahmad Wijaya", "role": "Vendor", "unit": "PT Maju Jaya"},
    "admin": {"password": "admin123", "nama": "Admin INAPROC", "role": "Administrator", "unit": "LKPP"}
}

def check_login(username: str, password: str) -> bool:
    if username in USER_CREDENTIALS:
        return USER_CREDENTIALS[username]["password"] == password
    return False

def get_user_profile(username: str) -> dict:
    return USER_CREDENTIALS.get(username)

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.current_user = None
    st.session_state.user_profile = None
    st.session_state.messages = []
    st.session_state.active_po = None
    st.session_state.vectorstore = None

if not st.session_state.logged_in:
    st.title("🔐 Login Chatbot INAPROC")
    col1, col2 = st.columns(2)
    with col1:
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        if st.button("Login", type="primary"):
            if check_login(username, password):
                st.session_state.logged_in = True
                st.session_state.current_user = username
                st.session_state.user_profile = get_user_profile(username)
                st.session_state.messages = []
                st.rerun()
            else:
                st.error("❌ Username atau password salah!")
    with col2:
        st.markdown("**Akun Demo:**\n- ppk_budi / 12345\n- staff_siti / 12345\n- vendor_ahmad / 12345\n- admin / admin123")
    st.stop()

@st.cache_resource
def build_knowledge_base_from_docx():
    """
    Membangun knowledge base dari file FAQ - Katalog Elektronik versi 6.docx
    """
    with st.spinner("📚 Memuat pengetahuan dari FAQ Katalog Elektronik versi 6..."):
        all_documents = []
        
        file_paths = [
            "FAQ - Katalog Elektronik versi 6.docx",
            "./FAQ - Katalog Elektronik versi 6.docx",
            "data/FAQ - Katalog Elektronik versi 6.docx",
        ]
        
        doc_path = None
        for path in file_paths:
            if os.path.exists(path):
                doc_path = path
                break
        
        if not doc_path:
            st.warning("⚠️ File FAQ tidak ditemukan. Gunakan pengetahuan default.")
            fallback = "INAPROC adalah platform pengadaan digital LKPP. Katalog Elektronik versi 6 adalah versi terbaru."
            doc = Document(page_content=fallback, metadata={"source": "fallback"})
            all_documents.append(doc)
        else:
            try:
                doc = docx.Document(doc_path)
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            full_text.append(" | ".join(row_text))
                
                doc_content = "\n".join(full_text)
                doc_chunk = Document(
                    page_content=doc_content,
                    metadata={"source": "FAQ - Katalog Elektronik versi 6.docx"}
                )
                all_documents.append(doc_chunk)
                st.success(f"✅ Berhasil membaca file: {doc_path}")
                
            except Exception as e:
                st.error(f"❌ Gagal membaca file: {str(e)}")
                fallback = "INAPROC adalah platform pengadaan digital LKPP."
                doc = Document(page_content=fallback, metadata={"source": "fallback"})
                all_documents.append(doc)
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1200,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = text_splitter.split_documents(all_documents)
        
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
        )
        vectorstore = FAISS.from_documents(chunks, embeddings)
        
        st.success(f"✅ Knowledge base siap! {len(chunks)} chunks")
        return vectorstore

if st.session_state.vectorstore is None:
    st.session_state.vectorstore = build_knowledge_base_from_docx()

@st.cache_data
def generate_orders():
    fake = Faker('id_ID')
    orders = {}
    for i in range(1, 31):
        po = f"PO-2025-{str(10000 + i)[1:]}"
        orders[po] = {
            "status": random.choice(["SHIPPED", "DELAYED", "DELIVERED"]),
            "kurir": random.choice(["JNE", "SiCepat"]),
            "resi": str(fake.random_number(digits=12)),
            "lokasi": fake.city(),
            "produk": random.choice(["Laptop ASUS", "AC Panasonic", "Proyektor Epson"]),
            "vendor": random.choice(["PT Maju Jaya", "CV Sukses Abadi"])
        }
    return orders

if "orders_db" not in st.session_state:
    st.session_state.orders_db = generate_orders()

def check_order_status(po_number: str) -> dict:
    if po_number in st.session_state.orders_db:
        return {"found": True, **st.session_state.orders_db[po_number]}
    return {"found": False}

def search_kb(query: str) -> str:
    """Cari di knowledge base, ambil 5 chunks teratas agar lebih lengkap"""
    try:
        docs = st.session_state.vectorstore.similarity_search(query, k=5)
        return "\n\n---\n\n".join([d.page_content for d in docs])
    except Exception as e:
        return ""

def get_bot_response(user_input: str) -> str:
    po_pattern = r'PO[-_]?\d{4,}[-_]?\d{3,}'
    po_match = re.search(po_pattern, user_input.upper())
    
    order_info = ""
    if po_match:
        po = po_match.group(0)
        st.session_state.active_po = po
        status = check_order_status(po)
        if status["found"]:
            order_info = f"Status: {status['status']}, Produk: {status['produk']}, Vendor: {status['vendor']}"
    
    kb_context = search_kb(user_input)
    
    profile = st.session_state.user_profile
    role = profile.get("role", "User") if profile else "User"
    
    history = ""
    if st.session_state.messages:
        for msg in st.session_state.messages[-4:]:
            history += f"{msg['role']}: {msg['content'][:100]}...\n"
    
    prompt = f"""
Anda adalah asisten chatbot resmi untuk INAPROC (Katalog Elektronik versi 6).

**PERAN ANDA:** Menjawab pertanyaan pengguna berdasarkan **FAQs - Katalog Elektronik versi 6** yang telah disediakan.

**ATURAN WAJIB (JANGAN DILANGGAR):**
1. ❌ JANGAN menambahkan informasi di luar dokumen FAQ!
2. ❌ JANGAN berhalusinasi atau membuat jawaban sendiri!
3. ✅ HANYA gunakan informasi dari bagian "INFORMASI DARI FAQ" di bawah ini.
4. ✅ Jika informasi tidak ada di FAQ, katakan "Maaf, informasi tersebut tidak tersedia di FAQ Katalog Elektronik versi 6. Silakan hubungi Pusat Bantuan INAPROC di layanan@lkpp.go.id atau WhatsApp 08111557709."
5. ✅ Jika ada di FAQ, jawab dengan detail dan tepat sesuai dokumen.
6. ✅ Sertakan nomor bagian/sub-bagian FAQ jika relevan (misal: "Berdasarkan FAQ 3.1.1.1.A...").

**ROLE USER:** {role}

**INFORMASI DARI FAQ (WAJIB DIPAKAI):**
{kb_context[:4000]}

**INFORMASI PESANAN (jika ada):**
{order_info}

**PERCAKAPAN SEBELUMNYA:**
{history}

**PERTANYAAN USER:** {user_input}

**JAWABAN (HARUS DARI FAQ, JANGAN HALUSINASI):**
"""
    
    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        return response.text
    except Exception as e:
        return f"⚠️ Error: {str(e)}"

profile = st.session_state.user_profile

col1, col2, col3 = st.columns([2, 2, 1])
with col1:
    st.title("🤖 Chatbot INAPROC")
    st.caption("Asisten resmi dengan pengetahuan FAQ Katalog Elektronik versi 6")
with col2:
    st.markdown(f"""
    <div class="user-profile-card">
        <strong>👤 {profile.get('nama', 'User') if profile else 'User'}</strong><br>
        <span class="role-badge">{profile.get('role', 'User') if profile else 'User'}</span>
    </div>
    """, unsafe_allow_html=True)
with col3:
    if st.button("🚪 Logout"):
        st.session_state.logged_in = False
        st.session_state.messages = []
        st.rerun()

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

if prompt := st.chat_input("Tanyakan status pesanan atau pertanyaan seputar INAPROC..."):
    with st.chat_message("user"):
        st.markdown(prompt)
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    with st.chat_message("assistant"):
        with st.spinner("Mencari di FAQ Katalog Elektronik..."):
            resp = get_bot_response(prompt)
            st.markdown(resp)
    st.session_state.messages.append({"role": "assistant", "content": resp})

with st.sidebar:
    st.markdown(f"### 👋 {profile.get('nama', 'User') if profile else 'User'}")
    if st.session_state.active_po:
        st.info(f"📌 PO Aktif: `{st.session_state.active_po}`")
    
    st.markdown("---")
    st.markdown("### 📚 Sumber Pengetahuan")
    st.markdown("- **FAQ Katalog Elektronik versi 6** (resmi LKPP)")
    st.markdown("- Panduan PPK/PP & Penyedia")
    
    st.markdown("---")
    if st.button("🗑️ Reset Chat"):
        st.session_state.messages = []
        st.session_state.active_po = None
        st.rerun()
    
    st.markdown("---")
    st.caption("🔐 Chatbot berbasis RAG dari dokumen resmi LKPP")
